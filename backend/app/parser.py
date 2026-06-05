import io
from pypdf import PdfReader
import docx

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extracts text from PDF bytes."""
    pdf_file = io.BytesIO(file_bytes)
    reader = PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        extracted = page.extract_text()
        if extracted:
            text += extracted + "\n"
    return text.strip()

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extracts text from DOCX bytes."""
    docx_file = io.BytesIO(file_bytes)
    doc = docx.Document(docx_file)
    text = []
    for paragraph in doc.paragraphs:
        if paragraph.text:
            text.append(paragraph.text)
    
    # Also extract text from tables if any
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    text.append(cell.text)
                    
    return "\n".join(text).strip()

def extract_text(file_bytes: bytes, filename: str) -> str:
    """Determines file type and extracts text accordingly."""
    lower_filename = filename.lower()
    if lower_filename.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif lower_filename.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif lower_filename.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="ignore").strip()
    else:
        raise ValueError("Unsupported file format. Please upload a PDF, DOCX, or TXT file.")
